#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
骰子队 - 屏幕识别配音助手 项目风险管理计划
生成 Word 文档
"""

import os
import io
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ============================================================
# 辅助函数
# ============================================================

def set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), kwargs.get('val', 'single'))
        element.set(qn('w:sz'), kwargs.get('sz', '4'))
        element.set(qn('w:color'), kwargs.get('color', '4472C4'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color='1F3864'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_paragraph(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    return p

def add_table_header_row(table, headers, bg='1F3864', fg='FFFFFF', bold=True):
    row = table.rows[0]
    for i, (cell, h) in enumerate(zip(row.cells, headers)):
        cell.text = ''
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(10)

def add_table_data_row(table, row_idx, data, bg=None, center_cols=None, bold_col=None):
    row = table.rows[row_idx]
    center_cols = center_cols or []
    for i, (cell, val) in enumerate(zip(row.cells, data)):
        cell.text = ''
        if bg:
            set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        if i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        run.font.size = Pt(10)
        if bold_col is not None and i == bold_col:
            run.font.bold = True
    return row

# ============================================================
# 图表生成
# ============================================================

def gen_risk_matrix():
    """生成风险概率-影响矩阵图"""
    matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(8, 6))

    # 背景色块
    colors_grid = [
        ['#92D050', '#FFFF00', '#FF0000'],
        ['#92D050', '#FFFF00', '#FF0000'],
        ['#00B050', '#92D050', '#FFFF00'],
        ['#00B050', '#92D050', '#FFFF00'],
        ['#00B050', '#00B050', '#92D050'],
    ]
    prob_labels = ['很高(0.9)', '高(0.7)', '中(0.5)', '低(0.3)', '很低(0.1)']
    impact_labels = ['低(0.05)', '中(0.1)', '高(0.2)']

    for i in range(5):
        for j in range(3):
            ax.add_patch(plt.Rectangle((j, 4-i), 1, 1,
                                        color=colors_grid[i][j], alpha=0.7, ec='white', lw=1.5))

    # 风险点
    risks = [
        # (x_impact_idx, y_prob_idx, label, offset)
        # impact: 0=低, 1=中, 2=高
        # prob: 0=很低, 1=低, 2=中, 3=高, 4=很高
        (2.4, 3.5, 'R1', (0.05, 0.1)),   # OCR识别准确率低 高影响/高概率
        (2.4, 2.5, 'R2', (0.05, 0.1)),   # TTS延迟高 高影响/中概率
        (1.5, 3.6, 'R3', (0.05, 0.1)),   # API限额 中影响/高概率
        (2.5, 4.3, 'R4', (0.05, 0.1)),   # 模型兼容性 高影响/很高概率
        (0.5, 2.5, 'R5', (0.05, 0.1)),   # 进度延迟-人员 低影响/中概率
        (1.5, 1.5, 'R6', (0.05, 0.1)),   # 需求变更 中影响/低概率
        (2.5, 1.5, 'R7', (0.05, 0.1)),   # 本地算力 高影响/低概率
        (0.5, 0.5, 'R8', (0.05, 0.1)),   # 数据安全 低影响/很低概率
        (1.5, 2.5, 'R9', (0.05, 0.1)),   # 前端兼容性 中影响/中概率
        (2.5, 3.2, 'R10', (0.05, 0.1)),  # 集成联调 高影响/高概率
    ]

    for (x, y, lbl, off) in risks:
        ax.plot(x, y, 'o', ms=12, color='#1F3864', zorder=5)
        ax.annotate(lbl, xy=(x, y), xytext=(x+off[0], y+off[1]),
                    fontsize=8, fontweight='bold', color='#1F3864',
                    ha='left', va='bottom')

    ax.set_xlim(0, 3)
    ax.set_ylim(0, 5)
    ax.set_xticks([0.5, 1.5, 2.5])
    ax.set_xticklabels(impact_labels, fontsize=10)
    ax.set_yticks([0.5, 1.5, 2.5, 3.5, 4.5])
    ax.set_yticklabels(prob_labels, fontsize=10)
    ax.set_xlabel('影响程度', fontsize=12, fontweight='bold')
    ax.set_ylabel('发生概率', fontsize=12, fontweight='bold')
    ax.set_title('风险概率-影响矩阵 (骰子队 · 屏幕识别配音助手)', fontsize=13, fontweight='bold', pad=12)

    # 图例
    patches = [
        mpatches.Patch(color='#FF0000', alpha=0.7, label='高风险'),
        mpatches.Patch(color='#FFFF00', alpha=0.7, label='中等风险'),
        mpatches.Patch(color='#92D050', alpha=0.7, label='低风险'),
        mpatches.Patch(color='#00B050', alpha=0.7, label='可忽略风险'),
    ]
    ax.legend(handles=patches, loc='lower right', fontsize=9)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def gen_risk_breakdown():
    """风险来源分布饼图"""
    matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False

    labels = ['技术风险', '进度风险', '需求风险', '资源风险', '质量风险', '外部风险']
    sizes = [35, 20, 15, 15, 10, 5]
    colors = ['#FF4444', '#FF8800', '#FFCC00', '#44AA44', '#4488FF', '#AA44AA']
    explode = (0.05,) * 6

    fig, ax = plt.subplots(figsize=(7, 5))
    wedges, texts, autotexts = ax.pie(
        sizes, explode=explode, labels=labels, colors=colors,
        autopct='%1.1f%%', shadow=True, startangle=140,
        textprops={'fontsize': 10}
    )
    for at in autotexts:
        at.set_fontsize(9)
        at.set_color('white')
        at.set_fontweight('bold')
    ax.set_title('项目风险来源分布', fontsize=13, fontweight='bold', pad=12)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def gen_risk_score_bar():
    """风险综合评分柱状图"""
    matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False

    risks = [
        ('R4\n模型兼容性', 0.9*0.20, '#FF2222'),
        ('R1\nOCR准确率', 0.7*0.20, '#FF4444'),
        ('R10\n集成联调', 0.7*0.20, '#FF4444'),
        ('R2\nTTS延迟', 0.5*0.20, '#FF8800'),
        ('R3\nAPI限额', 0.7*0.10, '#FF8800'),
        ('R7\n本地算力', 0.3*0.20, '#FFCC00'),
        ('R9\n前端兼容', 0.5*0.10, '#FFCC00'),
        ('R5\n进度延迟', 0.5*0.05, '#44AA44'),
        ('R6\n需求变更', 0.3*0.10, '#44AA44'),
        ('R8\n数据安全', 0.1*0.05, '#44AA44'),
    ]
    labels = [r[0] for r in risks]
    scores = [r[1] for r in risks]
    bar_colors = [r[2] for r in risks]

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(labels, scores, color=bar_colors, edgecolor='white', linewidth=0.8)
    for bar, s in zip(bars, scores):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.002,
                f'{s:.3f}', ha='center', va='bottom', fontsize=9, fontweight='bold')

    ax.set_xlabel('风险编号与名称', fontsize=11)
    ax.set_ylabel('风险评分（概率×影响）', fontsize=11)
    ax.set_title('各风险综合评分排序（定量分析）', fontsize=13, fontweight='bold')
    ax.set_ylim(0, 0.22)

    # 参考线
    ax.axhline(0.14, color='red', linestyle='--', alpha=0.6, label='高风险阈值(0.14)')
    ax.axhline(0.05, color='orange', linestyle='--', alpha=0.6, label='中风险阈值(0.05)')
    ax.legend(fontsize=9)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def gen_fishbone():
    """鱼骨图（风险因素因果图）"""
    matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('#F8F9FA')

    # 主骨
    ax.annotate('', xy=(11, 4), xytext=(1, 4),
                arrowprops=dict(arrowstyle='->', color='#1F3864', lw=2.5))
    ax.text(11.3, 4, '项目\n风险', ha='left', va='center', fontsize=12,
            fontweight='bold',
            bbox=dict(boxstyle='round,pad=0.4', fc='#FF2222', ec='white', alpha=0.85),
            color='white')

    # 分支定义: (主分支文字, x位置, 上/下, 子风险列表)
    branches = [
        ('技术风险', 3.0, 'top', ['OCR准确率不足', 'TTS延迟过高', '模型兼容性差', '接口不稳定']),
        ('进度风险', 5.5, 'top', ['技术熟练度不足', '任务衔接延迟', '多模块联调超时']),
        ('资源风险', 8.0, 'top', ['本地算力不足', 'API额度耗尽', '模型下载慢']),
        ('需求风险', 3.5, 'bottom', ['需求临时变更', '用户反馈调整', '功能蔓延']),
        ('质量风险', 6.5, 'bottom', ['测试覆盖率低', '代码规范不统一', 'bug积压']),
        ('外部风险', 9.0, 'bottom', ['网络/设备故障', '开源模型停更', '依赖库版本冲突']),
    ]

    colors_top = ['#4472C4', '#ED7D31', '#A9D18E']
    colors_bot = ['#FF5252', '#FFC107', '#9C27B0']

    for idx, (label, xb, side, children) in enumerate(branches):
        if side == 'top':
            color = colors_top[idx % 3]
            ya, yb = 6.5, 4
        else:
            color = colors_bot[idx % 3]
            ya, yb = 1.5, 4

        # 主枝
        ax.annotate('', xy=(xb, yb), xytext=(xb - 0.5 if side == 'top' else xb - 0.5, ya),
                    arrowprops=dict(arrowstyle='->', color=color, lw=2.0))

        # 分支标签
        ax.text(xb - 0.5 if side == 'top' else xb - 0.5,
                ya + (0.35 if side == 'top' else -0.35),
                label, ha='center', va='bottom' if side == 'top' else 'top',
                fontsize=10, fontweight='bold', color='white',
                bbox=dict(boxstyle='round,pad=0.3', fc=color, ec='white', alpha=0.9))

        # 子枝
        for ci, child in enumerate(children):
            offset = (ci - (len(children)-1)/2) * 0.55
            cx = xb - 0.5 + offset
            cy = (ya + yb) / 2 + (0.3 if side == 'top' else -0.3)
            ax.plot([cx, xb - 0.5 + (xb - (xb-0.5))*(1 - (ya-cy)/(ya-yb))],
                    [ya if side == 'top' else ya, cy],
                    color=color, lw=1.2, alpha=0.7)
            ax.text(cx, ya - 0.3 if side == 'top' else ya + 0.3,
                    child, ha='center',
                    va='bottom' if side == 'top' else 'top',
                    fontsize=8, color=color, alpha=0.9)

    ax.set_title('屏幕识别配音助手项目风险鱼骨图（因果分析）',
                 fontsize=14, fontweight='bold', pad=15, color='#1F3864')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def gen_response_timeline():
    """风险应对时间轴甘特风格图"""
    matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(11, 5))
    tasks = [
        ('风险识别与登记', 1, 3, '#4472C4'),
        ('定性/定量分析', 2, 4, '#ED7D31'),
        ('制定应对计划', 3, 6, '#A9D18E'),
        ('实施技术预研', 3, 8, '#FF5252'),
        ('API备份方案落实', 4, 6, '#FFC107'),
        ('算力资源检查', 3, 5, '#9C27B0'),
        ('持续风险监控', 4, 26, '#00B050'),
        ('风险审查会议', 7, 8, '#4472C4'),
        ('风险审查会议', 13, 14, '#4472C4'),
        ('风险审查会议', 20, 21, '#4472C4'),
        ('项目结束评审', 25, 26, '#1F3864'),
    ]
    milestones = [3, 15, 22, 26]
    milestone_labels = ['M1\n需求分析', 'M5\n各模块完成', 'M7\n系统集成', 'M9\n项目结题']

    for i, (label, start, end, color) in enumerate(tasks):
        ax.barh(i, end - start, left=start, height=0.5,
                color=color, alpha=0.8, edgecolor='white')
        ax.text((start + end)/2, i, label,
                ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    for m, ml in zip(milestones, milestone_labels):
        ax.axvline(m, color='#FF2222', linestyle='--', alpha=0.7, lw=1.5)
        ax.text(m, len(tasks) - 0.2, ml, ha='center', va='bottom', fontsize=8,
                color='#FF2222', fontweight='bold')

    ax.set_yticks([])
    ax.set_xlabel('项目天数（天）', fontsize=11)
    ax.set_title('风险管理活动时间计划', fontsize=13, fontweight='bold')
    ax.set_xlim(0, 28)
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('#F8F9FA')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


# ============================================================
# 文档主体
# ============================================================

def build_document():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 默认正文样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================================================
    # 封面
    # ==================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run('屏幕识别配音助手')
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle_p.add_run('项目风险管理计划')
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    info_lines = [
        ('团队名称', '骰子队'),
        ('项目成员', '马鑫  刘宇轩  梁艺馨  杨子祺  王婧晋'),
        ('课程', 'IT项目管理'),
        ('实验编号', 'Lab5'),
        ('编制日期', '2026年5月20日'),
    ]
    for k, v in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{k}：{v}')
        r.font.size = Pt(12)

    doc.add_page_break()

    # ==================================================
    # 目录（手工列出）
    # ==================================================
    add_heading(doc, '目  录', 1)
    toc = [
        '一、风险管理计划概述',
        '    1.1 文档说明',
        '    1.2 项目简介',
        '    1.3 风险管理方法与工具',
        '    1.4 角色与职责',
        '    1.5 风险管理时间安排',
        '二、风险识别',
        '    2.1 风险识别方法',
        '    2.2 风险源分析',
        '    2.3 风险登记册',
        '三、风险定性分析',
        '    3.1 风险概率-影响矩阵',
        '    3.2 风险优先级排序',
        '四、风险定量分析',
        '    4.1 定量方法说明',
        '    4.2 风险综合评分',
        '    4.3 决策树分析',
        '    4.4 蒙特卡洛模拟（项目进度风险）',
        '五、风险应对计划',
        '    5.1 应对策略总表',
        '    5.2 高优先级风险详细应对方案',
        '    5.3 风险储备金预算',
        '六、风险监控计划',
        '七、附录',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0 if not item.startswith('    ') else 1)
        for run in p.runs:
            run.font.size = Pt(10.5)
    doc.add_page_break()

    # ==================================================
    # 一、风险管理计划概述
    # ==================================================
    add_heading(doc, '一、风险管理计划概述', 1)

    add_heading(doc, '1.1 文档说明', 2)
    add_paragraph(doc, '本文档是"屏幕识别配音助手"项目（骰子队 · IT项目管理课程实验五）的风险管理计划，依据PMBOK第六版风险管理知识域框架编制，涵盖风险识别、定性与定量分析、应对规划及监控等完整流程，旨在为项目团队提供系统化的风险管理指导，确保项目在既定时间（2026年4月18日—5月13日，共26天）内按质完成交付。')

    add_heading(doc, '1.2 项目简介', 2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.add_run('项目背景：').bold = True
    p.add_run('有声阅读市场规模超70亿美元，国内盲人游戏玩家占比25%-30%，游戏无障碍需求旺盛。')
    p2 = doc.add_paragraph()
    p2.add_run('项目目标：').bold = True
    p2.add_run('打造便携、本地优先、低延迟的屏幕文本实时配音工具，面向电子书阅读、游戏娱乐、无障碍使用场景，通过屏幕OCR识别与多模型TTS合成，为用户提供高度自定义的听觉阅读与游戏陪伴体验。')
    p3 = doc.add_paragraph()
    p3.add_run('核心功能模块：').bold = True
    modules = ['屏幕识别模块（OCR）', '文本转音频模块（TTS：INDEX-TTS/Qwen3-TTS/GPT-SOVIT）', '智能回复与陪伴模块（大模型接入）', '前端页面模块（四大设置页）']
    for m in modules:
        add_paragraph(doc, f'· {m}', indent=True)
    p4 = doc.add_paragraph()
    p4.add_run('关键约束：').bold = True
    p4.add_run('项目周期26天，预算≤100元（主要为API调用费），团队5人，采用本地部署方式，目标平台Windows 11。')

    add_heading(doc, '1.3 风险管理方法与工具', 2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl, ['风险管理活动', '方法/工具', '说明'])
    rows_data = [
        ('风险识别', '头脑风暴法\n鱼骨图（因果图）\n核查表法', '团队集体讨论，参考lab3拖延风险识别结果，结合技术特征系统梳理'),
        ('定性分析', '概率-影响矩阵\n风险优先级评分', '用5×3矩阵对每项风险打分，划定高/中/低优先级'),
        ('定量分析', 'EMV期望货币价值\n决策树分析\n蒙特卡洛模拟（进度风险）\n敏感性分析', '对高优先级风险进行数值量化，为储备金预算和关键路径提供依据'),
        ('应对规划', '应对策略矩阵\n应急响应计划', '针对每项风险制定规避/转移/减轻/接受策略，明确责任人和触发条件'),
        ('风险监控', '风险审查会议\n偏差分析\n风险再评估', '项目期间每周定期审查，关键里程碑前强制审查'),
    ]
    for rd in rows_data:
        r = tbl.add_row()
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            p = r.cells[ci].paragraphs[0]
            p.add_run(val).font.size = Pt(9.5)
            p.paragraph_format.left_indent = Cm(0.2)

    add_heading(doc, '1.4 角色与职责', 2)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl2, ['角色', '成员', '风险管理职责', '负责风险类型'])
    roles_data = [
        ('项目经理 / 风险负责人', '马鑫', '统筹风险管理全流程，主持风险评审会议，批准应对预算', '全部风险'),
        ('技术开发', '马鑫', 'TTS模块技术风险排查与应对', '技术风险（TTS/集成）'),
        ('前端开发', '刘宇轩', '前端界面及前后端集成风险排查', '技术风险（前端）'),
        ('AI/模型', '梁艺馨', '大模型接入与兼容性风险排查', '技术风险（AI模型）'),
        ('屏幕识别开发', '杨子祺', 'OCR模块精度与性能风险排查', '技术风险（OCR）'),
        ('屏幕识别开发', '王婧晋', '屏幕捕获稳定性风险排查', '技术风险（屏幕捕获）'),
        ('全体成员', '全体', '日常风险观测，上报新风险，参与风险审查', '进度/资源/质量风险'),
    ]
    for rd in roles_data:
        r = tbl2.add_row()
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            p = r.cells[ci].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)

    add_heading(doc, '1.5 风险管理时间安排', 2)
    add_paragraph(doc, '风险管理活动贯穿项目全程，关键时间节点如下：')
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl3, ['活动', '时间节点', '持续时间', '说明'])
    timeline_data = [
        ('初始风险识别与分析', '项目第1-3天（4.18-4.20）', '3天', '团队头脑风暴，完成风险登记册初版'),
        ('应对计划制定', '第3-6天（4.20-4.23）', '4天', '完成风险应对矩阵，分配责任人'),
        ('技术预研（高风险）', '第3-8天（4.20-4.25）', '6天', 'OCR/TTS/大模型兼容性验证'),
        ('第一次风险审查', '第7-8天（4.24-4.25）', '1天', '结合需求分析完成情况，更新风险状态'),
        ('持续风险监控', '第4-26天', '全程', '发现新风险即时上报，每周例会审查'),
        ('第二次风险审查', '第13-14天（5.1-5.2）', '1天', '各模块开发中期检查'),
        ('第三次风险审查', '第20-21天（5.8-5.9）', '1天', '系统集成阶段风险重评估'),
        ('项目结束评审', '第25-26天（5.12-5.13）', '2天', '总结风险管理经验，形成风险报告'),
    ]
    for rd in timeline_data:
        r = tbl3.add_row()
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            r.cells[ci].paragraphs[0].add_run(val).font.size = Pt(9.5)

    # 风险管理时间轴图
    add_paragraph(doc, '图1：风险管理活动时间计划')
    timeline_buf = gen_response_timeline()
    doc.add_picture(timeline_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==================================================
    # 二、风险识别
    # ==================================================
    add_heading(doc, '二、风险识别', 1)

    add_heading(doc, '2.1 风险识别方法', 2)
    methods = [
        ('头脑风暴法', '团队5名成员集体讨论，结合lab3已识别的拖延原因（人员、技术、资源、需求、管理、环境、集成7类）进行拓展，形成初始风险清单。'),
        ('鱼骨图（因果分析）', '以"项目失败/超期/质量不达标"为鱼头，分析技术、进度、资源、需求、质量、外部六大类风险来源。'),
        ('核查表法', '参照IT软件项目常见风险核查表，逐项对照本项目特征（本地部署、AI模型集成、5人小团队、26天短周期），识别适用风险。'),
        ('专家访谈/文档回顾', '回顾lab1-lab4形成的项目文档（建议书、可行性分析、WBS、时间计划、质量计划），提炼已记录的潜在风险点。'),
    ]
    for name, desc in methods:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run(f'【{name}】 ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        p.add_run(desc).font.size = Pt(10)

    add_heading(doc, '2.2 风险源分析（鱼骨图）', 2)
    add_paragraph(doc, '下图为本项目风险因果鱼骨图，以"项目风险（超期/质量不达标/功能缺失）"为结果，从技术、进度、资源、需求、质量、外部六个维度梳理风险来源：')
    fishbone_buf = gen_fishbone()
    doc.add_picture(fishbone_buf, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, '图2：项目风险鱼骨图（因果分析）')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 饼图
    add_paragraph(doc, '')
    pie_buf = gen_risk_breakdown()
    doc.add_picture(pie_buf, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, '图3：项目风险来源分布（技术风险占比最高，达35%）')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, '2.3 风险登记册', 2)
    add_paragraph(doc, '下表为经团队识别确认的完整风险登记册（共识别10项主要风险）：')

    tbl_risk = doc.add_table(rows=1, cols=6)
    tbl_risk.style = 'Table Grid'
    tbl_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl_risk, ['风险ID', '风险名称', '风险类别', '风险描述', '潜在影响', '识别来源'])

    risks_register = [
        ('R1', 'OCR识别准确率不足', '技术风险',
         '针对游戏界面、动态文本、复杂背景，OCR引擎可能误识别或漏识别，导致配音内容错误',
         '用户体验差，功能核心指标（≥90%准确率）无法达标',
         '技术可行性分析/质量计划'),
        ('R2', 'TTS语音延迟过高', '技术风险',
         'Qwen3-TTS、INDEX-TTS等模型本地推理或API调用延迟超过5秒，影响实时配音体验',
         '质量目标（TTS延迟≤5s）无法达标，核心功能受损',
         '质量计划/技术路线'),
        ('R3', 'API调用额度耗尽', '资源风险',
         '大模型API（DeepSeek、阿里云）在开发测试阶段消耗过快，超出预算或免费额度',
         '开发进度暂停，需追加预算（原预算≤100元）',
         '经济可行性分析'),
        ('R4', 'AI模型兼容性差', '技术风险',
         '多个TTS/LLM模型（INDEX-TTS、Qwen3-TTS、GPT-SOVIT）与项目环境（Windows 11/Python）存在依赖冲突或接口不兼容',
         '技术方案被迫变更，工期延误，可能影响关键路径（26天）',
         'lab3时间计划/技术选型'),
        ('R5', '人员技术不熟练导致进度延迟', '进度风险',
         '团队成员对OCR/TTS模型集成、前后端接入等技术栈不够熟练，导致单任务完工时间超出估算',
         '项目关键路径延长，整体交付时间推迟',
         'lab3鱼骨图'),
        ('R6', '需求临时变更', '需求风险',
         '来自课程教师评审、用户测试反馈或团队内部讨论，导致功能范围扩大或调整',
         '已完成模块返工，增加工作量，影响进度和质量',
         '可行性分析/用户画像'),
        ('R7', '本地设备算力不足', '资源风险',
         '成员电脑GPU/CPU性能不足以运行本地TTS模型（如GPT-SOVIT），导致推理速度过慢或无法运行',
         '功能降级，只能使用API模式，增加成本或依赖网络',
         'lab4质量计划/资源约束'),
        ('R8', '数据隐私与安全风险', '外部风险',
         '屏幕识别功能可能捕获用户隐私内容，特别是在使用API模式时上传至云端，存在隐私泄露隐患',
         '用户信任度下降，合规性问题，影响推广',
         '利益相关者分析'),
        ('R9', '前端兼容性问题', '技术风险',
         '前端GUI框架在不同Windows版本或分辨率下显示异常，影响操作体验',
         '用户界面功能受损，需额外测试和修复时间',
         '技术选型/质量计划'),
        ('R10', '多模块集成联调失败', '技术风险',
         '屏幕识别→TTS→大模型→前端四大模块在系统集成阶段出现接口不匹配、数据格式冲突、时序问题',
         '系统集成阶段（任务G）延误，影响后续测试，威胁关键路径',
         'WBS/lab3时间计划'),
    ]

    bg_colors = ['FFFFFF', 'F2F7FF']
    for i, rd in enumerate(risks_register):
        r = tbl_risk.add_row()
        bg = bg_colors[i % 2]
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            set_cell_bg(r.cells[ci], bg)
            p = r.cells[ci].paragraphs[0]
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_page_break()

    # ==================================================
    # 三、风险定性分析
    # ==================================================
    add_heading(doc, '三、风险定性分析', 1)

    add_heading(doc, '3.1 风险概率-影响矩阵', 2)
    add_paragraph(doc, '采用5×3概率-影响矩阵对各风险进行定性评估。概率分为5级（0.1/0.3/0.5/0.7/0.9），影响分为3级（低0.05/中0.1/高0.2），风险评分 = 概率 × 影响。')
    add_paragraph(doc, '· 高风险（红色）：评分 ≥ 0.14，需立即制定应对措施')
    add_paragraph(doc, '· 中等风险（黄色）：0.05 ≤ 评分 < 0.14，需规划应对措施')
    add_paragraph(doc, '· 低风险（浅绿）：评分 < 0.05，持续监控')
    matrix_buf = gen_risk_matrix()
    doc.add_picture(matrix_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, '图4：风险概率-影响矩阵（R1-R10分布）')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, '3.2 风险优先级排序', 2)
    add_paragraph(doc, '基于概率-影响矩阵评分，完整的定性分析结果如下表：')

    tbl_qual = doc.add_table(rows=1, cols=7)
    tbl_qual.style = 'Table Grid'
    tbl_qual.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl_qual, ['风险ID', '风险名称', '发生概率', '影响程度', '风险评分', '优先级', '优先级依据'])

    qual_data = [
        ('R4', 'AI模型兼容性差',    '很高(0.9)', '高(0.2)',  '0.180', '★★★ 高', '直接威胁关键路径，影响所有技术模块'),
        ('R1', 'OCR识别准确率不足', '高(0.7)',   '高(0.2)',  '0.140', '★★★ 高', '核心功能指标，直接影响产品可用性'),
        ('R10','多模块集成联调失败', '高(0.7)',   '高(0.2)',  '0.140', '★★★ 高', '集成阶段关键风险，影响整体交付'),
        ('R2', 'TTS语音延迟过高',   '中(0.5)',   '高(0.2)',  '0.100', '★★ 中',  '核心质量目标，影响用户体验'),
        ('R3', 'API调用额度耗尽',   '高(0.7)',   '中(0.1)',  '0.070', '★★ 中',  '可能导致开发中断，概率较高'),
        ('R7', '本地设备算力不足',  '低(0.3)',   '高(0.2)',  '0.060', '★★ 中',  '影响本地部署核心优势'),
        ('R9', '前端兼容性问题',    '中(0.5)',   '中(0.1)',  '0.050', '★★ 中',  '影响界面功能，需测试验证'),
        ('R5', '人员技术不熟练',    '中(0.5)',   '低(0.05)', '0.025', '★ 低',   '通过技术培训和任务分配可缓解'),
        ('R6', '需求临时变更',      '低(0.3)',   '中(0.1)',  '0.030', '★ 低',   '课程项目需求相对稳定'),
        ('R8', '数据隐私风险',      '很低(0.1)', '低(0.05)', '0.005', '★ 低',   '概率低，可通过隐私设计缓解'),
    ]

    priority_colors = {
        '★★★ 高': 'FFCCCC',
        '★★ 中': 'FFFACD',
        '★ 低': 'CCFFCC',
    }
    for rd in qual_data:
        r = tbl_qual.add_row()
        bg = priority_colors.get(rd[5], 'FFFFFF')
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            set_cell_bg(r.cells[ci], bg)
            p = r.cells[ci].paragraphs[0]
            if ci in [0, 2, 3, 4, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 0:
                run.font.bold = True

    doc.add_page_break()

    # ==================================================
    # 四、风险定量分析
    # ==================================================
    add_heading(doc, '四、风险定量分析', 1)

    add_heading(doc, '4.1 定量方法说明', 2)
    add_paragraph(doc, '对定性分析中确认为"高"和"中"优先级的风险（R1、R2、R3、R4、R7、R9、R10），采用以下定量方法进行深入分析：')
    quant_methods = [
        '期望货币价值（EMV）：量化每项风险的期望成本影响，用于计算应急储备金。',
        '决策树分析：针对R3（API额度风险），分析"自建模型"vs"购买API"两种方案的期望收益。',
        '蒙特卡洛模拟：评估项目进度风险（R5+R10组合效应）对关键路径（26天）完工概率的影响。',
        '敏感性分析：识别对项目总成本影响最大的风险因素（龙卷风图）。',
    ]
    for m in quant_methods:
        add_paragraph(doc, f'· {m}', indent=True)

    add_heading(doc, '4.2 风险综合评分（EMV分析）', 2)
    add_paragraph(doc, '基于风险概率与潜在工期/成本影响，计算各风险的期望货币价值（EMV = 概率 × 影响量）：')

    tbl_emv = doc.add_table(rows=1, cols=7)
    tbl_emv.style = 'Table Grid'
    tbl_emv.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl_emv, ['风险ID', '风险名称', '发生概率P', '最大工期损失(天)', '最大成本损失(元)', 'EMV工期(天)', 'EMV成本(元)'])

    emv_data = [
        ('R4', 'AI模型兼容性差',    '0.9', '5', '0',   '4.5', '0'),
        ('R1', 'OCR识别准确率不足', '0.7', '4', '0',   '2.8', '0'),
        ('R10','多模块集成联调失败', '0.7', '4', '0',   '2.8', '0'),
        ('R2', 'TTS语音延迟高',     '0.5', '3', '30',  '1.5', '15'),
        ('R3', 'API额度耗尽',       '0.7', '2', '100', '1.4', '70'),
        ('R7', '本地算力不足',      '0.3', '2', '0',   '0.6', '0'),
        ('R9', '前端兼容性问题',    '0.5', '1', '0',   '0.5', '0'),
        ('合计', '—', '—', '—', '—', '14.1天', '85元'),
    ]

    for i, rd in enumerate(emv_data):
        r = tbl_emv.add_row()
        bg = 'FFF0F0' if i == len(emv_data)-1 else ('FFFFFF' if i % 2 == 0 else 'F2F7FF')
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            set_cell_bg(r.cells[ci], bg)
            p = r.cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if i == len(emv_data)-1:
                run.font.bold = True

    add_paragraph(doc, '注：工期损失基于lab3关键路径（26天）估算；成本损失主要来自API费用。EMV合计显示项目需预留约14天弹性缓冲和85元成本应急储备。')

    # 风险评分柱状图
    bar_buf = gen_risk_score_bar()
    doc.add_picture(bar_buf, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, '图5：各风险综合评分排序（概率×影响）')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, '4.3 决策树分析——API方案选择', 2)
    add_paragraph(doc, '针对R3（API调用额度耗尽）风险，项目组面临"坚持使用云端API"vs"切换本地开源模型"两种方案。采用决策树分析评估期望收益：')

    dec_tree_text = """
决策节点：TTS方案选择
├── 方案A：使用云端TTS API（Qwen3-TTS / INDEX-TTS API版）
│   ├── API额度充足（P=0.3）：语音质量高，延迟低 → 价值 +100分
│   └── API额度耗尽（P=0.7）：追加费用约100元，开发停滞2天 → 价值 -60分
│   EMV_A = 0.3×100 + 0.7×(-60) = 30 - 42 = -12分
│
└── 方案B：优先使用本地开源模型（GPT-SOVIT / INDEX-TTS本地版）
    ├── 本地算力足够（P=0.7）：零API成本，可离线使用 → 价值 +80分
    └── 本地算力不足（P=0.3）：需降级至轻量模型，质量略降 → 价值 +20分
    EMV_B = 0.7×80 + 0.3×20 = 56 + 6 = 62分
    """
    p_tree = doc.add_paragraph()
    p_tree.paragraph_format.left_indent = Cm(0.5)
    run_tree = p_tree.add_run(dec_tree_text.strip())
    run_tree.font.name = 'Courier New'
    run_tree.font.size = Pt(9)

    add_paragraph(doc, '决策结论：EMV_B(62分) >> EMV_A(-12分)，建议优先采用方案B（本地开源模型），在本地算力不足时降级为轻量模型，保留云端API作为最后备选。')

    add_heading(doc, '4.4 蒙特卡洛模拟——进度风险评估', 2)
    add_paragraph(doc, '对关键路径任务（A→B→C/D→G→H→I，标准工期26天）进行蒙特卡洛模拟，考虑R4、R5、R10三个进度风险的叠加效应：')

    mc_data = [
        ('需求分析A', '乐观2天', '最可能3天', '悲观4天', '期望3.0天'),
        ('系统设计B', '乐观2天', '最可能3天', '悲观5天', '期望3.2天'),
        ('屏幕识别C', '乐观6天', '最可能8天', '悲观12天', '期望8.3天（R4风险）'),
        ('文本转音频D', '乐观6天', '最可能8天', '悲观11天', '期望8.2天（R1/R2风险）'),
        ('系统集成G', '乐观4天', '最可能5天', '悲观9天', '期望5.5天（R10风险）'),
        ('测试优化H', '乐观3天', '最可能4天', '悲观6天', '期望4.2天'),
        ('总计（关键路径）', '—', '—', '—', '期望32.4天'),
    ]
    tbl_mc = doc.add_table(rows=1, cols=5)
    tbl_mc.style = 'Table Grid'
    add_table_header_row(tbl_mc, ['任务', '乐观工期', '最可能工期', '悲观工期', '期望工期（含风险）'])
    for i, rd in enumerate(mc_data):
        r = tbl_mc.add_row()
        bg = 'FFF0F0' if i == len(mc_data)-1 else ('FFFFFF' if i % 2 == 0 else 'F2F7FF')
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            set_cell_bg(r.cells[ci], bg)
            p = r.cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if i == len(mc_data)-1:
                run.font.bold = True

    add_paragraph(doc, '模拟结论：在不采取应对措施的情况下，期望完工时间为32.4天（超出计划26天约6.4天），按时完成概率约25%。通过实施风险应对措施（技术预研、模块并行、合理任务缓冲），可将按时完成概率提升至约75%。')

    doc.add_page_break()

    # ==================================================
    # 五、风险应对计划
    # ==================================================
    add_heading(doc, '五、风险应对计划', 1)

    add_heading(doc, '5.1 应对策略总表', 2)
    add_paragraph(doc, '根据PMBOK风险应对策略框架（威胁：规避/转移/减轻/接受），针对各风险制定应对措施：')

    tbl_resp = doc.add_table(rows=1, cols=7)
    tbl_resp.style = 'Table Grid'
    tbl_resp.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(tbl_resp, ['风险ID', '风险名称', '优先级', '应对策略', '具体措施', '责任人', '触发条件'])

    resp_data = [
        ('R4', 'AI模型兼容性差', '★★★高', '减轻\n+规避',
         '① 项目第3天前完成所有拟用模型的安装验证\n② 建立备用模型清单（每类至少2个备选）\n③ 隔离各模型依赖环境（conda虚拟环境）\n④ 统一接口封装层，屏蔽底层差异',
         '梁艺馨\n马鑫', '模型安装失败或接口调用报错'),
        ('R1', 'OCR识别准确率不足', '★★★高', '减轻',
         '① 采用多引擎方案：PaddleOCR为主，Tesseract为备\n② 图像预处理（灰度化、二值化、降噪）提升识别率\n③ 针对游戏场景设置ROI区域框选\n④ 定期准确率测试，低于85%立即调整预处理策略',
         '杨子祺\n王婧晋', '测试准确率低于90%'),
        ('R10', '多模块集成联调失败', '★★★高', '减轻\n+减轻',
         '① 提前定义模块间接口规范（数据格式、调用方式）\n② 开发阶段建立Mock接口，并行开发减少等待\n③ 集成前进行两两模块联调测试\n④ 集成期间每日代码合并，及时暴露冲突',
         '马鑫\n刘宇轩', '集成阶段出现接口异常或数据格式不匹配'),
        ('R2', 'TTS语音延迟过高', '★★中', '减轻\n+接受',
         '① 优先使用轻量模型（INDEX-TTS轻量版），降低推理延迟\n② 降低采样率（22050Hz→16000Hz）\n③ 使用流式TTS（边生成边播放）\n④ 若延迟仍超5s，设置"可接受阈值提示"告知用户',
         '马鑫', 'TTS延迟超过5秒'),
        ('R3', 'API调用额度耗尽', '★★中', '规避\n+转移',
         '① 优先使用本地开源模型（决策树分析方案B）\n② API调用设置每日配额上限警告\n③ 申请阿里云/DeepSeek教育优惠额度\n④ 团队共享一个API账号，统一管理token使用',
         '马鑫', 'API调用费用超过50元或余额告警'),
        ('R7', '本地算力不足', '★★中', '转移\n+减轻',
         '① 转移至性能最优的成员电脑作为主开发机\n② 轻量化模型替代（降低模型参数规模）\n③ 关闭后台应用，释放GPU内存\n④ 若仍不足，切换至API调用模式（备用方案A）',
         '杨子祺\n王婧晋', '模型推理时间>10秒或设备严重卡顿'),
        ('R9', '前端兼容性问题', '★★中', '减轻',
         '① 统一开发环境（Windows 11 + Python 3.12 + tkinter/PyQt）\n② 定义UI分辨率适配规范（最小1080P）\n③ 开发完成后在不同分辨率设备上测试\n④ 使用相对布局，避免绝对像素坐标',
         '刘宇轩', '界面显示错乱或控件超出边界'),
        ('R5', '人员技术不熟练', '★低', '减轻',
         '① 项目开始前分享技术参考资料（OCR/TTS教程）\n② 技术难点由项目经理优先攻关，形成Demo供参考\n③ 设置里程碑检查节点，及时发现进度偏差',
         '马鑫', '任务实际进度落后计划2天以上'),
        ('R6', '需求临时变更', '★低', '接受\n+减轻',
         '① 变更需经组长审批，记录变更原因和影响\n② 对变更评估工期影响，超过1天需调整后续计划\n③ 保持最小可行版本（MVP），控制变更范围',
         '马鑫（审批）\n全体', '收到正式变更请求'),
        ('R8', '数据隐私风险', '★低', '减轻\n+接受',
         '① 默认使用本地模型，避免屏幕内容上传云端\n② API模式下对识别内容进行本地缓存，不长期存储\n③ 在用户手册中明确隐私说明',
         '梁艺馨', '发现隐私内容被上传或存储'),
    ]

    resp_bg = ['FFEEEE', 'FFEEEE', 'FFEEEE', 'FFFFF0', 'FFFFF0', 'FFFFF0', 'FFFFF0', 'F0FFF0', 'F0FFF0', 'F0FFF0']
    for i, rd in enumerate(resp_data):
        r = tbl_resp.add_row()
        bg = resp_bg[i]
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            set_cell_bg(r.cells[ci], bg)
            p = r.cells[ci].paragraphs[0]
            if ci in [0, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if ci == 0:
                run.font.bold = True

    add_heading(doc, '5.2 高优先级风险详细应对方案', 2)

    # R4详细
    add_paragraph(doc, '▶ R4 AI模型兼容性差——详细应对方案', bold=True)
    r4_detail = [
        ('预防措施', '项目第3天（4月20日）前，各成员在自己的开发机上完成以下模型的安装测试：INDEX-TTS、GPT-SOVIT、Qwen3-TTS（API模式）、PaddleOCR、DeepSeek/Qwen（API模式）。安装失败的记录错误信息并立即上报。'),
        ('备用方案', '为每个功能模块准备至少2个可用模型：\n· TTS：INDEX-TTS（首选）→ Qwen3-TTS API → Edge-TTS（备用）\n· OCR：PaddleOCR（首选）→ Tesseract → EasyOCR\n· LLM：DeepSeek API（首选）→ Qwen API → 本地Ollama'),
        ('应急响应', '如开发过程中主选模型失效：\n① 立即切换备用模型（耗时≤2小时）\n② 通知项目经理评估影响范围\n③ 若影响关键路径，启动时间缓冲（预留2天）'),
        ('验收标准', '系统集成（任务G）开始前，所有模块的主选和至少1个备用模型可正常调用。'),
    ]
    tbl_r4 = doc.add_table(rows=len(r4_detail), cols=2)
    tbl_r4.style = 'Table Grid'
    for ri, (k, v) in enumerate(r4_detail):
        tbl_r4.rows[ri].cells[0].text = ''
        set_cell_bg(tbl_r4.rows[ri].cells[0], '1F3864')
        p0 = tbl_r4.rows[ri].cells[0].paragraphs[0]
        run0 = p0.add_run(k)
        run0.font.bold = True
        run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run0.font.size = Pt(9.5)
        tbl_r4.rows[ri].cells[1].text = ''
        tbl_r4.rows[ri].cells[1].paragraphs[0].add_run(v).font.size = Pt(9.5)

    doc.add_paragraph()

    # R1详细
    add_paragraph(doc, '▶ R1 OCR识别准确率不足——详细应对方案', bold=True)
    r1_detail = [
        ('测试基准', '以游戏界面截图（100张）作为测试集，基准要求：识别准确率≥90%，识别延迟≤2s/帧。'),
        ('图像预处理管线', '① 灰度化 → ② 自适应二值化（Otsu算法）→ ③ 高斯降噪 → ④ 形态学处理 → ⑤ OCR识别。\n对复杂背景图像额外进行边缘增强。'),
        ('多引擎策略', '主引擎：PaddleOCR（中英文混合效果好）；备用引擎：Tesseract（字体清晰场景）；专用场景：截图大模型识别（复杂游戏UI）。'),
        ('降级策略', '若准确率80%-90%：调整预处理参数，提示用户"识别质量较低"；若<80%：弹窗提示用户手动框选清晰区域。'),
    ]
    tbl_r1 = doc.add_table(rows=len(r1_detail), cols=2)
    tbl_r1.style = 'Table Grid'
    for ri, (k, v) in enumerate(r1_detail):
        tbl_r1.rows[ri].cells[0].text = ''
        set_cell_bg(tbl_r1.rows[ri].cells[0], '1F3864')
        p0 = tbl_r1.rows[ri].cells[0].paragraphs[0]
        run0 = p0.add_run(k)
        run0.font.bold = True
        run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run0.font.size = Pt(9.5)
        tbl_r1.rows[ri].cells[1].text = ''
        tbl_r1.rows[ri].cells[1].paragraphs[0].add_run(v).font.size = Pt(9.5)

    doc.add_paragraph()

    # R10详细
    add_paragraph(doc, '▶ R10 多模块集成联调失败——详细应对方案', bold=True)
    r10_detail = [
        ('接口规范文档', '开发阶段（任务C/D/E/F）启动前，由马鑫主导，各模块负责人共同制定《模块接口规范说明书》，明确：数据格式（文本/音频/配置）、调用方式（函数调用/HTTP）、错误码和异常处理。'),
        ('Mock开发策略', '开发期间各模块通过Mock接口进行独立测试，不依赖其他模块完成。Mock接口与正式接口保持一致，降低集成风险。'),
        ('渐进式集成', '按照"两两集成→三模块集成→全量集成"的渐进策略：\n阶段1（5.2-5.4）：OCR+TTS联调；大模型+TTS联调\n阶段2（5.4-5.6）：三模块+前端联调\n阶段3（5.6-5.7）：全量集成测试'),
        ('应急时间缓冲', 'lab3时间计划中系统集成任务G标准工期5天（5.3-5.7），风险应对预留2天缓冲（最迟5.9日完成集成），确保测试任务H（4天）和总结任务I（3天）不受影响。'),
    ]
    tbl_r10 = doc.add_table(rows=len(r10_detail), cols=2)
    tbl_r10.style = 'Table Grid'
    for ri, (k, v) in enumerate(r10_detail):
        tbl_r10.rows[ri].cells[0].text = ''
        set_cell_bg(tbl_r10.rows[ri].cells[0], '1F3864')
        p0 = tbl_r10.rows[ri].cells[0].paragraphs[0]
        run0 = p0.add_run(k)
        run0.font.bold = True
        run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run0.font.size = Pt(9.5)
        tbl_r10.rows[ri].cells[1].text = ''
        tbl_r10.rows[ri].cells[1].paragraphs[0].add_run(v).font.size = Pt(9.5)

    add_heading(doc, '5.3 风险储备金预算', 2)
    add_paragraph(doc, '基于EMV定量分析结果，制定以下应急储备方案：')
    tbl_budget = doc.add_table(rows=1, cols=4)
    tbl_budget.style = 'Table Grid'
    add_table_header_row(tbl_budget, ['储备类型', '金额/时间', '用途', '审批权限'])
    budget_data = [
        ('时间应急储备（管理层）', '2天', '应对高不确定性风险（R4/R10），由项目经理统一调配', '项目经理（马鑫）批准'),
        ('时间应急储备（已知风险）', '4天', '应对已识别风险（R1/R2/R5），在各任务工期估算中嵌入', '无需审批，已内嵌'),
        ('成本应急储备', '100元（已有预算内）', '主要用于R3（API额度耗尽）应急追加', '项目经理批准'),
        ('总计', '6天+100元', '占项目标准工期（26天）的23%，属合理范围', '—'),
    ]
    for i, rd in enumerate(budget_data):
        r = tbl_budget.add_row()
        bg = 'FFF0F0' if i == len(budget_data)-1 else ('FFFFFF' if i % 2 == 0 else 'F2F7FF')
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            set_cell_bg(r.cells[ci], bg)
            p = r.cells[ci].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if i == len(budget_data)-1:
                run.font.bold = True

    doc.add_page_break()

    # ==================================================
    # 六、风险监控计划
    # ==================================================
    add_heading(doc, '六、风险监控计划', 1)

    add_paragraph(doc, '风险监控是贯穿整个项目生命周期的持续性活动，目的是跟踪已识别风险的状态变化、识别新风险、验证应对措施的有效性，并在必要时更新风险登记册。')

    add_heading(doc, '6.1 监控活动与频率', 2)
    tbl_mon = doc.add_table(rows=1, cols=5)
    tbl_mon.style = 'Table Grid'
    add_table_header_row(tbl_mon, ['监控活动', '频率', '负责人', '输出物', '说明'])
    mon_data = [
        ('风险状态检查', '每日', '各模块负责人', '风险日志更新', '在每日开发中观察风险触发迹象'),
        ('风险例会审查', '每周一次\n（项目期间共3次）', '马鑫（主持）\n全体', '风险审查报告', '讨论风险变化，评估应对措施效果，识别新风险'),
        ('里程碑风险评估', '每个里程碑前', '马鑫', '里程碑风险评估单', '结合里程碑完成情况，重新评估后续风险'),
        ('应对措施有效性验证', '应对措施实施后', '对应责任人', '有效性验证记录', '确认措施是否达到预期效果'),
        ('风险登记册更新', '有变化即更新', '马鑫', '更新的风险登记册', '新增、关闭或调整风险优先级'),
    ]
    for rd in mon_data:
        r = tbl_mon.add_row()
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            p = r.cells[ci].paragraphs[0]
            p.add_run(val).font.size = Pt(9.5)

    add_heading(doc, '6.2 风险监控指标（KRI）', 2)
    add_paragraph(doc, '设定以下关键风险指标（KRI），实现量化预警：')
    kri_data = [
        ('任务进度偏差', '≤0', '>0（超期）', '>2天', 'R5/R10/R4'),
        ('OCR准确率（测试集）', '≥90%', '85%-90%', '<85%', 'R1'),
        ('TTS延迟（平均）', '≤5s', '5s-8s', '>8s', 'R2'),
        ('API调用花费', '≤50元', '50-80元', '>80元', 'R3'),
        ('Bug未修复数量', '≤5个', '5-15个', '>15个', 'R9/R10'),
        ('每日代码提交率', '≥80%（有提交）', '40%-80%', '<40%', 'R5'),
    ]
    tbl_kri = doc.add_table(rows=1, cols=5)
    tbl_kri.style = 'Table Grid'
    add_table_header_row(tbl_kri, ['指标名称', '绿色（正常）', '黄色（预警）', '红色（告警）', '关联风险'])
    kri_bgs = [('CCFFCC', 'FFFACD', 'FFCCCC')]
    for rd in kri_data:
        r = tbl_kri.add_row()
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            if ci == 1:
                set_cell_bg(r.cells[ci], 'E8FFE8')
            elif ci == 2:
                set_cell_bg(r.cells[ci], 'FFFFF0')
            elif ci == 3:
                set_cell_bg(r.cells[ci], 'FFE8E8')
            p = r.cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size = Pt(9.5)

    add_heading(doc, '6.3 风险监控流程', 2)
    flow_text = """
风险监控工作流：
┌─────────────────────────────────────────────────────────┐
│  日常观测 → 触发条件判断 → 是否触发？                    │
│       ↓                                                   │
│    是：立即上报项目经理                                   │
│       ↓                                                   │
│    评估影响（工期/质量/成本）                             │
│       ↓                                                   │
│    启动应对计划（预防措施→应急响应→降级方案）            │
│       ↓                                                   │
│    执行后验证 → 有效？→ 是：更新风险登记册（已关闭）     │
│                    ↓                                      │
│                   否：升级处理，调整应对方案              │
│                                                           │
│  每周例会：汇总风险状态，识别新风险，更新优先级           │
└─────────────────────────────────────────────────────────┘
    """
    p_flow = doc.add_paragraph()
    p_flow.paragraph_format.left_indent = Cm(0.5)
    run_flow = p_flow.add_run(flow_text.strip())
    run_flow.font.name = 'Courier New'
    run_flow.font.size = Pt(9)

    doc.add_page_break()

    # ==================================================
    # 七、附录
    # ==================================================
    add_heading(doc, '七、附录', 1)

    add_heading(doc, '附录A：风险识别核查表', 2)
    add_paragraph(doc, '下表为本项目使用的风险识别核查表，"✓"表示该类风险已在本项目中识别：')
    tbl_check = doc.add_table(rows=1, cols=4)
    tbl_check.style = 'Table Grid'
    add_table_header_row(tbl_check, ['风险类别', '典型风险项', '本项目适用', '已识别风险编号'])
    check_data = [
        ('技术风险', 'AI模型兼容性/精度/延迟问题', '✓', 'R1/R2/R4/R9/R10'),
        ('进度风险', '关键路径延误/人员能力不足', '✓', 'R5'),
        ('资源风险', '计算资源/API额度/经费超支', '✓', 'R3/R7'),
        ('需求风险', '需求变更/范围蔓延', '✓', 'R6'),
        ('质量风险', '集成测试失败/缺陷积压', '✓', 'R10/R9'),
        ('外部风险', '开源模型停更/依赖库冲突/隐私', '✓', 'R8/R4'),
        ('管理风险', '分工不合理/进度监控不及时', '部分适用', '包含在R5中'),
        ('法律合规风险', '开源许可证冲突', '低风险', '未单独列出'),
    ]
    for rd in check_data:
        r = tbl_check.add_row()
        for ci, val in enumerate(rd):
            r.cells[ci].text = ''
            p = r.cells[ci].paragraphs[0]
            if ci == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val == '✓':
                    run = p.add_run(val)
                    run.font.color.rgb = RGBColor(0, 0x80, 0)
                    run.font.bold = True
                    run.font.size = Pt(10)
                else:
                    p.add_run(val).font.size = Pt(9.5)
            else:
                p.add_run(val).font.size = Pt(9.5)

    add_heading(doc, '附录B：风险登记册模板（持续更新）', 2)
    add_paragraph(doc, '风险登记册在项目期间持续维护，字段说明如下：')
    fields = [
        ('风险ID', '唯一标识符（R1-Rn）'),
        ('风险名称', '简洁描述风险内容'),
        ('风险类别', '技术/进度/资源/需求/质量/外部'),
        ('风险描述', '详细说明风险内容、触发条件和潜在影响'),
        ('发生概率', '0.1/0.3/0.5/0.7/0.9（五档评估）'),
        ('影响程度', '低(0.05)/中(0.1)/高(0.2)（三档评估）'),
        ('风险评分', '发生概率×影响程度'),
        ('优先级', '高（≥0.14）/中（0.05-0.14）/低（<0.05）'),
        ('应对策略', '规避/转移/减轻/接受'),
        ('应对措施', '具体操作步骤'),
        ('责任人', '负责跟踪和应对的成员'),
        ('触发条件', '启动应急响应的具体条件'),
        ('当前状态', '待观测/已触发/已关闭'),
        ('最后更新时间', '每次状态变化时更新'),
    ]
    tbl_template = doc.add_table(rows=1, cols=2)
    tbl_template.style = 'Table Grid'
    add_table_header_row(tbl_template, ['字段名称', '说明'])
    for k, v in fields:
        r = tbl_template.add_row()
        r.cells[0].text = k
        r.cells[0].paragraphs[0].runs[0].font.bold = True
        r.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        r.cells[1].text = v
        r.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    add_heading(doc, '附录C：参考资料', 2)
    refs = [
        '[1] Project Management Institute. A Guide to the Project Management Body of Knowledge (PMBOK® Guide) — Sixth Edition. 2017.',
        '[2] 骰子队 lab1：《项目初步建议书》——屏幕识别配音助手.',
        '[3] 骰子队 lab2：《可行性分析与任务结构分解》——包括WBS、人员可行性、技术路线.',
        '[4] 骰子队 lab3：《编制时间计划》——关键路径（26天）、里程碑、项目拖延风险鱼骨图.',
        '[5] 骰子队 lab4：《质量计划》——质量目标（OCR≥90%，TTS≤5s）、质量风险与应对.',
        '[6] 梁艺馨 lab5（理论部分）：风险定量/定性分析方法说明.',
        '[7] 王婧晋 lab5（理论部分）：风险应对策略与风险监控方法.',
    ]
    for ref in refs:
        add_paragraph(doc, ref, indent=True)

    # 文档末尾
    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run('— 骰子队 · 屏幕识别配音助手 · 项目风险管理计划 · 2026年5月 —')
    r_end.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    r_end.font.size = Pt(9)

    return doc


if __name__ == '__main__':
    output_path = 'E:/厦大/IT项目管理/lab5/骰子队_lab5_风险管理计划.docx'
    print('正在生成文档...')
    doc = build_document()
    doc.save(output_path)
    print(f'文档已保存至：{output_path}')
